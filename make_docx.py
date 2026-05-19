from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 스타일 설정 ──────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def set_font(run, size=10, bold=False, color=None, name='맑은 고딕'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(31, 73, 125))
    # 하단 선
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(68, 114, 196))
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(91, 155, 213))
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    for line in text.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # 헤더
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        # 헤더 배경색
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6F1')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('B2B AI Agent Platform')
set_font(run, size=28, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('포트폴리오 기술 문서')
set_font(run, size=14, color=(68, 114, 196))

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('생산기업 · 물류기업을 연결하는 AI Agent 납품 플랫폼')
set_font(run, size=11, color=(100, 100, 100))

doc.add_page_break()

# ════════════════════════════════════════════════════
# 목차
# ════════════════════════════════════════════════════
heading1('목  차')
toc_items = [
    ('1.', '프로젝트 개요'),
    ('2.', '시스템 아키텍처'),
    ('3.', '데이터베이스 설계'),
    ('4.', 'AI Agent 기능'),
    ('5.', '기술 스택'),
    ('6.', '주요 화면'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'  {num}  {title}')
    set_font(run, size=11)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 1. 프로젝트 개요
# ════════════════════════════════════════════════════
heading1('1. 프로젝트 개요')

heading2('1.1 기획 배경')
body('AI Agent를 납품하는 플랫폼 기업의 관점에서, 제조기업(EM)과 물류기업을 연결하는 B2B 통합 시스템을 구현하였다.')
body('생산기업은 수주·생산·재고를 관리하고, 물류기업은 배차·운송을 담당한다. 플랫폼은 양사 데이터를 조용히 수집하며 중간에서 자동 연동을 수행한다.')

heading2('1.2 핵심 목표')
bullet('B/L(선하증권) 업로드 → 물류 배차 요청 자동 생성')
bullet('생산 완료(READY_TO_SHIP) → 출고 배차 요청 자동 생성')
bullet('물류 귀로 최적화: 부산→서울 납품 후 서울→부산 화물 자동 매칭')
bullet('플랫폼이 양사 데이터를 모니터링하며 기업 간 조율')

heading2('1.3 비즈니스 모델')
body('플랫폼 기업은 AI Agent를 각 기업에 납품하고, 모든 거래 데이터를 축적한다.')
body('신규 생산기업 가입 시 기존 물류 인프라를 즉시 제공할 수 있으며, 데이터 기반 물류 최적화가 가능하다.')

# ════════════════════════════════════════════════════
# 2. 시스템 아키텍처
# ════════════════════════════════════════════════════
heading1('2. 시스템 아키텍처')

heading2('2.1 3개 기업 구조')
add_table(
    ['서비스', '백엔드 포트', '프론트 포트', '역할'],
    [
        ['플랫폼 AGENT', '3000', '5175', '통합 모니터링 + 기업 간 자동 연동'],
        ['생산 AGENT', '3001', '5173', '수주·생산·재고·B/L 관리'],
        ['물류 AGENT', '3002', '5174', '배차·차량·운송 관리'],
        ['MySQL 8.0', '3306', '-', '중앙 데이터 저장소'],
    ],
    col_widths=[4, 3, 3, 7]
)

heading2('2.2 데이터 흐름')
code_block('''
[생산기업] B/L PDF 업로드
    → OpenAI GPT-4o-mini 파싱
    → em_agent.bl_imports 저장
    → logistics_agent.dispatch_requests 자동 생성 (IMPORT)

[생산기업] 수주 READY_TO_SHIP 전환
    → logistics_agent.dispatch_requests 자동 생성 (PRODUCTION)

[플랫폼] 양사 DB 크로스 쿼리 모니터링
    → 배달 완료 차량 감지 → 귀로 화물 자동 매칭 (개발 중)
''')

heading2('2.3 실행 환경')
bullet('단일 npm start 명령으로 3개 서버 동시 실행 (concurrently)')
bullet('Docker Compose로 MySQL 환경 표준화')
bullet('init/ SQL 파일 자동 실행으로 DB 즉시 재현 가능')

# ════════════════════════════════════════════════════
# 3. 데이터베이스 설계
# ════════════════════════════════════════════════════
heading1('3. 데이터베이스 설계')

heading2('3.1 DB 구성')
body('MySQL 8.0 기반 3개 독립 데이터베이스로 기업별 데이터를 분리하여 관리한다.')
add_table(
    ['데이터베이스', '소유 기업', '역할'],
    [
        ['em_agent', '생산기업 (EM)', '수주, 생산공정, 재고, B/L 관리'],
        ['logistics_agent', '물류기업', '차량, 배차 요청, 운송 이력'],
        ['platform_agent', '플랫폼 기업', '등록 기업 관리, Agent 행동 로그'],
    ],
    col_widths=[4, 4, 9]
)

heading2('3.2 em_agent 테이블 정의')
add_table(
    ['테이블명', '주요 컬럼', '설명'],
    [
        ['orders', 'id, order_number, product_type, quantity, status, due_date', '수주 목록 (이스트우드 발주 20건)'],
        ['raw_stock_receipts', 'material_type, quantity, received_date, bl_id', '원자재 입고 내역 (B/L 연결)'],
        ['bl_imports', 'bl_number, vessel, origin_port, eta, pickup_address', 'B/L 선하증권 등록'],
        ['material_specs', 'material_type, weight_kg, unit_qty', '소재별 무게 기준'],
        ['first_process_sessions', 'order_id, target_qty, output_qty, status', '1차 공정 작업 이력'],
        ['second_process_sessions', 'order_id, machine_id, units_per_hour, status', '2차 공정 작업 이력'],
        ['event_logs', 'event_type, qty, material_type, timestamp', '전체 이벤트 로그'],
    ],
    col_widths=[4, 7.5, 5.5]
)

heading2('3.3 logistics_agent 테이블 정의')
add_table(
    ['테이블명', '주요 컬럼', '설명'],
    [
        ['vehicles', 'plate_number, vehicle_type, capacity_kg, base_region, driver_name', '차량 10대 (거점지역 포함)'],
        ['dispatch_requests', 'request_type, cargo_desc, weight_kg, pickup_location, em_receipt_id, em_order_id', '배차 요청 (EM 자동 연동)'],
        ['dispatches', 'request_id, vehicle_id, assigned_at, estimated_delivery, status', '실제 배차 내역'],
        ['event_logs', 'event_type, dispatch_id, vehicle_id, timestamp', '배차 이벤트 로그'],
    ],
    col_widths=[4, 8, 5]
)

heading2('3.4 기업 간 연동 방식')
body('물류기업은 생산기업 DB를 직접 조회하지 않는다. 모든 연동은 이벤트 기반으로 처리된다.')

heading3('B/L 등록 → 수입 배차 자동 생성')
bullet('EM 서버: POST /api/bl-imports 요청 처리 시 logistics_agent.dispatch_requests INSERT')
bullet('em_receipt_id (UNIQUE) 로 중복 방지')
bullet('무게는 material_specs 기준으로 서버에서 자동 산출')

heading3('생산 완료 → 출고 배차 자동 생성')
bullet('EM 서버: PATCH /api/orders/:id 에서 status → READY_TO_SHIP 감지')
bullet('logistics_agent.dispatch_requests에 PRODUCTION 타입으로 자동 INSERT')
bullet('em_order_id (UNIQUE) 로 중복 방지')

heading3('플랫폼 크로스 DB 쿼리')
bullet('플랫폼 서버는 em_agent.*, logistics_agent.* 직접 쿼리 가능 (동일 MySQL 인스턴스)')
bullet('모니터링 전용으로만 사용, 데이터 변경은 각 기업 서버를 통해서만 처리')

heading2('3.5 주요 외래키 관계')
add_table(
    ['테이블', '컬럼', '참조', '용도'],
    [
        ['raw_stock_receipts', 'bl_id', 'bl_imports.id', 'B/L과 입고 내역 연결'],
        ['dispatch_requests', 'em_receipt_id', 'em_agent.raw_stock_receipts.id', '수입 입고건 중복 방지'],
        ['dispatch_requests', 'em_order_id', 'em_agent.orders.id', '출고건 중복 방지'],
        ['dispatches', 'request_id', 'dispatch_requests.id', '배차 요청과 실행 연결'],
        ['dispatches', 'vehicle_id', 'vehicles.id', '배차된 차량 연결'],
    ],
    col_widths=[4, 4, 5.5, 4]
)

heading2('3.6 시드 데이터')
add_table(
    ['항목', '내용'],
    [
        ['수주', '이스트우드 발주 20건 / 총 716,000장 / 납기 2026-06-30'],
        ['원자재 입고', '19건 (B/L 연결 포함)'],
        ['소재 무게 기준', 'FABRIC 1.2kg/롤 / STICKER_PAPER 1.3kg/롤 / CHIP 1.5kg/5만개'],
        ['차량', '10대 (1톤 2대·5톤 4대·11톤 2대·25톤 2대 / 서울·인천·경기·부산 거점)'],
        ['B/L', '6건 (원단·RFID칩·스티커지 수입 건)'],
    ],
    col_widths=[5, 12]
)

heading2('3.7 Docker 환경 표준화')
body('MySQL Docker 이미지의 /docker-entrypoint-initdb.d/ 기능을 활용하여 init/ 폴더의 SQL 파일이 컨테이너 최초 실행 시 자동으로 순서대로 실행된다.')
code_block('''
docker compose up -d   # MySQL 실행 + DB/테이블/시드 데이터 자동 구성
npm install
npm start              # 3개 서버 동시 실행
''')
body('이를 통해 개발 환경에 관계없이 동일한 DB 상태를 즉시 재현할 수 있다.')

# ════════════════════════════════════════════════════
# 4. AI Agent 기능
# ════════════════════════════════════════════════════
heading1('4. AI Agent 기능')

heading2('4.1 B/L PDF 자동 파싱')
bullet('사용자가 B/L PDF 업로드 시 서버가 pdf-parse로 텍스트 추출')
bullet('OpenAI GPT-4o-mini에 전달 → 화물명·수량·ETA·항구 정보 JSON 자동 추출')
bullet('파싱 결과를 확인 후 입고 등록 → 물류 배차 요청 자동 생성')

heading2('4.2 수입 입고 → 물류 배차 자동 연동')
bullet('B/L 확정 시 em_agent 서버가 logistics_agent.dispatch_requests에 직접 INSERT')
bullet('IMPORT 타입, 소재별 무게 자동 계산, 중복 등록 방지')

heading2('4.3 생산 완료 → 출고 배차 자동 연동')
bullet('수주 상태가 READY_TO_SHIP으로 변경되는 순간 물류 배차 요청 자동 생성')
bullet('PRODUCTION 타입으로 픽업(공장) → 배달(고객사) 자동 설정')

heading2('4.4 AI 자동 배차')
bullet('물류기업 대시보드에서 AI 배차 버튼 클릭 시 Gemini 2.5 Flash 호출')
bullet('화물 무게, 차량 적재량, 기사 거점 지역 기반 최적 차량 선택')
bullet('API 할당량 초과 시 무게 기준 폴백 자동 처리')

heading2('4.5 재고 AI 조언')
bullet('현재 재고 / 입고 예정 일정 / 전체 수주 필요량 / 1차 공정 완료 재고 분석')
bullet('Gemini 2.5 Flash 기반 한국어 조언 생성 (SAFE / CHECK 상태 판단)')
bullet('수주별 일일 생산 필요량 자동 계산 (납기 역산)')

heading2('4.6 귀로 최적화 (개발 예정)')
bullet('부산→서울 납품 완료 후 배달 완료 차량을 플랫폼이 감지')
bullet('서울→부산 방향 출고 화물 자동 매칭 → 빈차 복귀 없이 화물 배정')
bullet('각 기업 단독으로는 불가능한 매칭 — 플랫폼만 양사 데이터 보유')

# ════════════════════════════════════════════════════
# 5. 기술 스택
# ════════════════════════════════════════════════════
heading1('5. 기술 스택')
add_table(
    ['구분', '기술', '용도'],
    [
        ['Frontend', 'React 19 + Vite', 'UI 구성'],
        ['Frontend', 'Tailwind CSS', '스타일링'],
        ['Backend', 'Node.js + Express', 'REST API 서버 (ESM 모듈)'],
        ['Database', 'MySQL 8.0', '중앙 데이터 저장소'],
        ['Database', 'mysql2/promise', 'Node.js MySQL 드라이버 (커넥션 풀)'],
        ['AI', 'OpenAI GPT-4o-mini', 'B/L PDF 파싱, 채팅'],
        ['AI', 'Google Gemini 2.5 Flash', 'AI 배차, 재고 진단'],
        ['DevOps', 'Docker Compose', 'MySQL 실행 환경 표준화'],
        ['DevOps', 'concurrently', '3개 서버 단일 명령 실행'],
    ],
    col_widths=[3, 5, 9]
)

# ════════════════════════════════════════════════════
# 6. 주요 화면
# ════════════════════════════════════════════════════
heading1('6. 주요 화면')
body('※ 시연 화면은 별도 첨부 또는 GitHub 저장소 시연 사진 폴더 참조')
body('GitHub: https://github.com/smw312500-commit/B2B-PLATFORM')

bullet('생산 AGENT — 수주 목록, 생산 공정, 재고 현황, B/L 업로드')
bullet('물류 AGENT — 배차 대기 목록, AI 자동 배차, 차량 관리')
bullet('플랫폼 AGENT — 양사 통합 모니터링 대시보드')

# ── 저장 ──────────────────────────────────────────────────────
output_path = r'e:\PROJECT\B2B Platform\B2B_AI_Agent_Platform_포트폴리오.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
